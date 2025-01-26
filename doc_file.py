from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from erfandb import read_data
from msg_handling import MsgHandling
import flet as ft


def save_doc(page: ft.Page, start_date, end_date):
    document = Document()

    heading = document.add_heading("<لیست بیماران عرفان>", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # define error function for showing error message
    error_msg = MsgHandling(page)

    try:

        row_len, id, name, date, khadamat, prices = list(read_data(page, start_date, end_date))
        # remove "," from price
        new_price = []
        for price in prices:
            if "," in price:
                new_price.append(int(price.replace(",", "")))

        table = document.add_table(rows=row_len + 2, cols=4, style="Table Grid")
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def table_text(x, y, text):

            cell = table.cell(x, y)
            cell.text = text

            # Apply formatting: Center horizontally and vertically
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Set font size
            cell.paragraphs[0].runs[0].font.size = Pt(14)

        def save_data(y, list_text):
            for row, text in enumerate(list_text):
                cell = table.cell(row + 1, y)
                cell.text = str(text)

                # Apply formatting: Center horizontally and vertically
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Set font size
                cell.paragraphs[0].runs[0].font.size = Pt(14)

        # سرتیتر جدول تعریف شده
        table_text(0, 3, "نام بیماران")
        table_text(0, 2, "تاریخ پذیرش")
        table_text(0, 1, "خدمات انجام شده")
        table_text(0, 0, "قیمت")

        # ذخیره اطلاعات در ورد
        save_data(3, name)
        save_data(2, date)
        save_data(1, khadamat)
        save_data(0, prices)

        # تعداد بیمار
        table_text(len(name) + 1, 3, "تعداد بیمار")
        table_text(len(name) + 1, 2, str(len(name)))

        # قیمت خدمات ارائه شده
        table_text(len(name) + 1, 1, "مبلغ کل خدمات")
        table_text(len(name) + 1, 0, format(f"{sum(new_price):,}"))

        document.save("erfan.doc")

    except ValueError:
        # define error dialog when we want to choose data err.args[0]
        error_msg.error_dialog(ValueError.args,
                               400, 100, 20, False, ft.TextAlign.LEFT)
